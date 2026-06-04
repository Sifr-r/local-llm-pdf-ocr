"""
docx_writer - Utility for exporting OCR text to Word Documents (.docx).
"""

from __future__ import annotations

import io
import re
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Inches, Pt

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph


def convert_markdown_to_docx(markdown_text: str) -> io.BytesIO:
    """
    Parse a Markdown string and generate a polished Word Document (.docx) as a binary stream.

    Supports:
    - Page breaks and titles (from '## Page X' or '--- PAGE X ---')
    - Markdown headings: # (H1), ## (H2), ### (H3)
    - Bullet lists (- or *)
    - Numbered lists (1. , etc.)
    - Inline styles: bold (**text**), italic (*text*), bold+italic (***text***), code (`text`)
    - Paragraph spacing and font customisation (Arial, 11pt, 1.15 line spacing)
    """
    doc = Document()

    # Configure page setup: Standard 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configure base style: Arial 11pt
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # Simple Markdown Parser
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Page Breaks (e.g., "## Page X" or "--- PAGE X ---")
        is_page_hdr = stripped.startswith("## Page ") or (
            stripped.startswith("--- PAGE ") and stripped.endswith(" ---")
        )
        if is_page_hdr:
            if doc.paragraphs:  # Skip page break for the very first page
                doc.add_page_break()
            heading_text = stripped.replace("-", "").replace("#", "").strip()
            h = doc.add_heading(heading_text, level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Headings H1, H2, H3
        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        elif stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(4)

        # Bullet Lists
        elif stripped.startswith("- ") or stripped.startswith("* "):
            content = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            _add_inline_formatting(p, content)

        # Numbered Lists (matches "1. ", "12. ", etc.)
        elif re.match(r"^\d+\.\s+", stripped):
            match = re.match(r"^(\d+)\.\s+(.*)", stripped)
            if match:
                content = match.group(2)
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(3)
                _add_inline_formatting(p, content)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                _add_inline_formatting(p, stripped)

        # Normal Paragraphs
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            _add_inline_formatting(p, stripped)

        i += 1

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _add_inline_formatting(paragraph: Paragraph, text: str) -> None:
    """Helper to parse markdown bold/italic/code inline constructs and add runs to paragraph."""
    # Split text into formatted chunks and normal text
    pattern = re.compile(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)")
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue

        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(part)
